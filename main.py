#!/usr/bin/env python3
# coding=utf-8
import os
import sys
import xlsxwriter

from PIL import Image
from PyQt5 import QtGui
from PyQt5.QtGui import QPixmap
from PyQt5.QtWidgets import *
from PyQt5.uic import loadUi
from PyQt5 import QtCore
from docx import Document
from docx.shared import Inches

# Основной класс программы
class Main(QDialog):
    path_to_file = ""
    def __init__(self):
        super(Main, self).__init__()
        loadUi('Form1.ui', self)  # Загрузка формы из файла

        # Задание заголовка окна
        self.setWindowTitle('Разработка кроссплатформенных приложений')

        # Привязываем к кнопкам наши процедуры-обработчики
        self.pushButton_Run.clicked.connect(self.solve)
        self.pushButton_Img.clicked.connect(self.img)

    # Процедура решения примера
    def solve(self):
        array = []
        array.append(self.lineEdit_0.text())
        array.append(self.lineEdit_1.text())
        array.append(self.lineEdit_2.text())
        array.append(self.lineEdit_3.text())
        array.append(self.lineEdit_4.text())
        array.append(self.lineEdit_5.text())

        my_file = array[0] + '.xlsx'  # Имя файла
        book = xlsxwriter.Workbook(my_file)  # Создание файла
        sheet = book.add_worksheet()  # Добавление в него книги
        sheet.set_column('A:A', 20)  # Установка ширины колонки
        bold = book.add_format({'bold': True})  # Формат жирного текста
        cf = book.add_format()

        cf.set_pattern(1)  # This is optional when using a solid fill.
        cf.set_bg_color('blue')

        sb5 = book.add_format({'border': 5})
        sb1 = book.add_format({'border': 1})
        row = 6
        sheet.merge_range('A1:B1', array[2] + " - " + array[1] + " " + array[0], sb5)  # Выдача текста в ячейку

        sheet.write('A2', array[3], sb1)  # Выдача значения в ячейку 3 строка 1 столбец [2,0]
        sheet.write_string('B2', array[4] + ' год', cf)  # Выдача значения в ячейку 3 строка 1 столбец [2,0]
        sheet.write('A3', array[5], sb1)  # Выдача значения в ячейку 3 строка 1 столбец [2,0]
        if self.path_to_file != "":
            sheet.insert_image('B3', self.path_to_file, {'x_scale': 0.03, 'y_scale': 0.03})  # Вставка в ячейку картинки

        book.close()  # Закрытие файла
        os.system('start ' + array[0] + '.xlsx') # Открытие файла

        document = Document()
        document.add_heading(array[1] + " " + array[0], 0)
        p = document.add_paragraph(array[4] + ' год опыта в ' + array[2] + "\n")
        p.add_run('Мои сильные стороны:' + array[3]).bold = True
        document.add_paragraph('')
        if self.path_to_file != "":
            document.add_picture(self.path_to_file, width=Inches(5.25))
        document.add_heading('ПРОСЬБА НЕ ТРЕВОЖИТЬ!!!', level=1)
        document.add_paragraph(array[5]).bold = True
        document.save(array[0] + '.docx')

        os.system('start ' + array[0] + '.docx') # Открытие файла

    def img(self):

        self.path_to_file = QFileDialog.getOpenFileName(self, 'Открыть файл', '',
                                                   "Text Files (*.png)")[0]

        if self.path_to_file:
            f = Image.open(self.path_to_file, 'r')
            self.label_ImgName.setText(self.path_to_file)

        (width, height) = f.size
        if width > height:
            h = height * (251 / width)
            rez = QtCore.QSize(251, int(h))
        else:
            w = width * (141 / height)
            rez = QtCore.QSize(int(w), 141)
        pixmap = QtGui.QPixmap(self.path_to_file)

        pixmap = pixmap.scaled(rez)

        self.label_Img.setPixmap(pixmap)


# Основная часть программы
app = QApplication(sys.argv)
window = Main()  # базовый класс окна
window.show()  # отобразить окно на экране
sys.exit(app.exec_())  # запуск основного цикла приложения